from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def paragraph_record(paragraph):
    return {
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "runs": [
            {
                "text": run.text,
                "bold": bool(run.bold),
                "italic": bool(run.italic),
                "font": run.font.name,
                "size_pt": run.font.size.pt if run.font.size else None,
            }
            for run in paragraph.runs
        ],
    }


def inspect(path: Path):
    doc = Document(path)
    return {
        "file": str(path),
        "paragraphs": [paragraph_record(p) for p in doc.paragraphs if p.text.strip()],
        "tables": [
            [[cell.text for cell in row.cells] for row in table.rows]
            for table in doc.tables
        ],
        "sections": len(doc.sections),
        "headers": [
            [p.text for p in section.header.paragraphs if p.text.strip()]
            for section in doc.sections
        ],
        "footers": [
            [p.text for p in section.footer.paragraphs if p.text.strip()]
            for section in doc.sections
        ],
    }


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    target = Path(sys.argv[1])
    paths = [target] if target.is_file() else sorted(target.glob("*.docx"))
    print(json.dumps([inspect(path) for path in paths], ensure_ascii=False, indent=2))
