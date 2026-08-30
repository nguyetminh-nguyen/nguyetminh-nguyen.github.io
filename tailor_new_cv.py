from __future__ import annotations

import copy
import json
import shutil
import sys
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\nguye\OneDrive\Tài liệu\CV & Motivation Letters\New CV")
BACKUP_ROOT = Path(r"C:\website\cv_backups_before_tailoring")


COMMON = {
    "education": [
        "MSc Financial Economics - Erasmus University Rotterdam\t2025-2026",
        "Relevant coursework: Risk Management, Advanced Investments, Behavioral Finance\tRotterdam, NL",
        "BSc Economics & Business Economics (IBEB) - Erasmus University\t2022-2025",
        "GPA 8.3 (Cum Laude), Thesis 8.5\tRotterdam, NL",
    ],
}


CVS = {
    "Minh Nguyen CV - GCS.docx": {
        "profile": (
            "Financial Economics master's student who enjoys getting data, processes, and admin work into a clear and workable shape. "
            "I have built recurring reports, checks, and simple automation in Excel and Alteryx, and I am comfortable writing documentation, "
            "following up actions, and working with people across teams. I learn new tools quickly and like finding practical ways to make everyday work more reliable and efficient."
        ),
        "ideex": [
            "Improved recurring Excel and Alteryx workflows, cutting preparation time by about 40% and making updates more repeatable.",
            "Prepared recurring and ad hoc Excel analyses, comparison tables, and presentations, checking details before sharing findings with stakeholders.",
            "Worked with Payroll and Compensation to collect, validate, and map payroll information for EU Pay Transparency readiness.",
            "Learned Pay Analytics and ServiceNow quickly while coordinating requests, records, and follow-ups with colleagues across Total Rewards.",
        ],
        "ves": [
            "Built a repeatable Stata workflow to clean and merge VES 2021-2024 data into approximately 3.7 million firm-year records.",
            "Standardized and documented 181 data mappings across firm, investment, and energy files to keep definitions traceable across years.",
            "Checked duplicate keys, missing values, outliers, and structural changes before releasing data for analysis and reporting.",
        ],
        "projects": {
            "first_title": "Reporting & Quality-Control Pipeline - Excel, Power Query & Power BI",
            "first_bullets": [
                "Consolidated data from multiple sources into a repeatable Excel and Power Query workflow for recurring reporting.",
                "Built cleaning, validation, and consistency checks so issues could be identified before reports were shared.",
                "Created a Power BI dashboard that made Actual, Budget, and Forecast differences easier to review with stakeholders.",
                "Documented the update process and controls so the workflow could be repeated without rebuilding it from scratch.",
            ],
            "second_title": "FP&A Planning & Performance Model",
            "second_bullets": [
                "Built monthly, YTD, and full-year P&L views comparing Actual, Budget, and Forecast across four regions.",
                "Linked revenue, COGS, OPEX, and EBITDA variances to operational drivers across regions and periods.",
                "Compared Upside, Base, and Downside scenario impacts on FY2026 revenue, gross margin, and EBITDA by region.",
                "Built a Power BI dashboard combining executive P&L, variance, and scenario views for management review.",
            ],
        },
        "skills": [
            "Business & Project Support: Microsoft Excel, Word, PowerPoint, Meeting Notes, Documentation, Action Tracking",
            "Data & Process Improvement: Power Query, Power BI, Data Validation, Reporting, Workflow Improvement",
            "Tools: Alteryx, SQL, Python (pandas); comfortable learning unfamiliar systems and technology",
            "Working Style: Organised, detail-focused, proactive, clear communicator, reliable follow-through",
        ],
    },
    "Minh Nguyen CV - UL Solutions.docx": {
        "profile": (
            "Financial Economics master's student with hands-on experience checking financial and people data, reconciling inputs, and preparing clear Excel reports for stakeholders. "
            "I work carefully with details, learn new systems quickly, and like improving repetitive work through practical automation. Now looking to bring these strengths to customer-account, credit, and collections work in an international finance team."
        ),
        "ideex": [
            "Improved recurring Excel and Alteryx workflows, cutting preparation time by about 40% and reducing manual errors.",
            "Cleaned and validated raw data, built comparison tables, and prepared recurring and ad hoc Excel analyses for stakeholder reporting.",
            "Compared Korn Ferry and Willis Towers Watson market data across 20+ countries, checking consistency before use in salary-range reviews.",
            "Worked with Payroll to collect, validate, and map payroll elements across two countries, keeping the data and supporting records organised.",
        ],
        "ves": [
            "Built a repeatable Stata workflow to clean and merge VES 2021-2024 data into approximately 3.7 million firm-year records.",
            "Standardized and verified 181 data mappings across firm, investment, and energy files, preserving a clear audit trail.",
            "Detected duplicate keys, missing values, outliers, and structural changes across years to prevent misleading analysis.",
        ],
        "projects": {
            "first_title": "Accounting Close & Financial Reporting Model - Excel & Power Query",
            "first_bullets": [
                "Built a Power Query-driven close model using FY2025 general-ledger and subledger data, including AR and AP source data.",
                "Generated 120 month-end adjustment journals for depreciation, inventory, payroll accruals, and debt or lease interest.",
                "Produced a post-close trial balance across 93 accounts with total debits and credits of EUR 89.34 million and zero net imbalance.",
                "Created monthly P&L and year-end balance-sheet reports from reusable Power Query outputs and PivotTables.",
            ],
            "second_title": "Customer Profitability Analysis | Power BI",
            "second_bullets": [
                "Built a Power BI report tracking revenue, gross margin, ARR, NRR, and churn across customer segments.",
                "Combined customer-level revenue with cost-to-serve and acquisition-cost data to calculate profit after acquisition cost.",
                "Flagged loss-making customers and linked profitability gaps to service costs and underlying operational drivers.",
                "Created customer deep dives and management watchlists to highlight portfolio risks and priority actions.",
            ],
        },
        "skills": [
            "Finance Operations: Financial Data Validation, Reconciliation, Financial Reporting, Month-End Close Support",
            "Data & Reporting: Excel (Advanced), Power Query, Power BI, Alteryx, SQL, Python (pandas)",
            "Customer & Stakeholder Support: Clear Communication, Record Keeping, Issue Follow-Up, Documentation",
            "Working Style: Accurate, organised, calm with multiple priorities, proactive about process improvement",
        ],
    },
    "Minh Nguyen CV - Rosewood Amsterdam.docx": {
        "profile": (
            "Financial Economics master's student with practical experience keeping financial data accurate, building reconciled reports, and turning raw information into clear follow-up actions. "
            "I enjoy careful, hands-on finance work, learn new systems quickly, and look for simple ways to reduce manual effort without losing control of the details. Interested in growing in Accounts Receivable and financial operations in a client-focused environment."
        ),
        "ideex": [
            "Improved recurring Excel and Alteryx workflows, cutting preparation time by about 40% while reducing manual errors.",
            "Cleaned and validated raw data, maintained comparison tables, and prepared recurring Excel analyses and reports for stakeholders.",
            "Improved compensation recommendation templates and kept supporting market-data files and documentation organised for recurring use.",
            "Worked with Payroll and Compensation to collect and check inputs across teams, resolving gaps before reports were finalised.",
        ],
        "ves": [
            "Built a repeatable Stata workflow to clean and merge VES 2021-2024 data into approximately 3.7 million firm-year records.",
            "Standardized and verified 181 data mappings across firm, investment, and energy files to keep records consistent and traceable.",
            "Performed checks for duplicates, missing values, outliers, and structural changes before data was used for reporting.",
        ],
        "projects": {
            "first_title": "Accounting Close & Financial Reporting Model - Excel & Power Query",
            "first_bullets": [
                "Built a Power Query-driven close model using FY2025 general-ledger and subledger data, including AR and AP source data.",
                "Generated 120 month-end adjustment journals for depreciation, inventory, payroll accruals, and debt or lease interest.",
                "Produced a post-close trial balance across 93 accounts with total debits and credits of EUR 89.34 million and zero net imbalance.",
                "Created monthly P&L and year-end balance-sheet reports from reusable Power Query outputs and PivotTables.",
            ],
            "second_title": "Customer Profitability Analysis | Power BI",
            "second_bullets": [
                "Built a Power BI report tracking revenue, gross margin, ARR, NRR, and churn across customer segments.",
                "Combined customer-level revenue with cost-to-serve and acquisition-cost data to calculate profit after acquisition cost.",
                "Flagged loss-making customers and linked profitability gaps to service costs and underlying operational drivers.",
                "Created customer deep dives and management watchlists to highlight portfolio risks and priority actions.",
            ],
        },
        "skills": [
            "Financial Operations: Financial Data Validation, Reconciliation, Financial Reporting, Month-End Close Support",
            "Data & Tools: Excel (Advanced), Power Query, Power BI, Alteryx, SQL, Python (pandas)",
            "Administration: Documentation, Record Keeping, Cross-Functional Follow-Up, Clear Written Communication",
            "Working Style: Accurate, customer-minded, organised, independent, always looking for practical improvements",
        ],
    },
    "Minh Nguyen CV - McCain Foods.docx": {
        "profile": (
            "Analytical Financial Economics master's student who enjoys working with operational data and turning it into clear, useful reporting. I have built repeatable Excel, Power Query, and Power BI workflows, with strong attention to data accuracy and follow-through. I learn new systems quickly and like finding sensible automation that makes everyday processes more reliable and effective."
        ),
        "ideex": [
            "Improved recurring Excel and Alteryx workflows, cutting preparation time by about 40% and making updates more repeatable.",
            "Prepared and checked recurring and ad hoc Excel analyses, comparison tables, and reports for Total Rewards stakeholders.",
            "Collected, validated, and mapped payroll information with Payroll and Compensation teams for EU Pay Transparency readiness.",
            "Maintained an EU Pay Transparency tracker, incorporated weekly updates, and communicated relevant changes to the team.",
        ],
        "ves": [
            "Built a repeatable Stata workflow to clean and merge VES 2021-2024 data into approximately 3.7 million firm-year records.",
            "Standardized and verified 181 data mappings across firm, investment, and energy files, documenting definitions and comparability.",
            "Checked duplicates, missing values, outliers, and structural changes to keep operational data reliable for analysis.",
        ],
        "projects": {
            "first_title": "FP&A Planning & Performance Model",
            "first_bullets": [
                "Built a driver-based Excel model covering Actual, Budget, and multiple Forecast versions across four regions.",
                "Created linked P&L schedules from revenue through COGS, gross profit, payroll, non-payroll OPEX, and EBITDA.",
                "Added data-quality controls for Actual-to-Forecast cutovers, customer and FTE roll-forwards, scenario tie-outs, and P&L reconciliation.",
                "Built executive Power BI views linking P&L performance and operational KPIs to variance drivers and management commentary.",
            ],
            "second_title": "Procurement & Spend Analytics - Power BI",
            "second_bullets": [
                "Built a four-page Power BI report tracking spend, budget variance, realized and pipeline savings, contract coverage, and on-time delivery.",
                "Analyzed price leakage, off-contract, emergency, and maverick spend across categories, suppliers, channels, and business units.",
                "Created supplier-risk views combining delivery reliability, quality, cost, contract status, and single-source exposure.",
                "Used category and supplier drill-downs to explain whether adverse spend outcomes came from price, volume, compliance, mix, or concentration.",
            ],
        },
        "skills": [
            "Operational Reporting: Excel (Advanced), Power Query, Power BI, KPI Analysis, Data Validation, Reconciliation",
            "Analysis & Automation: Alteryx, SQL, Python (pandas), Repeatable Workflows, Quality Controls",
            "Business Support: Documentation, Clear Reporting, Cross-Functional Coordination, Structured Problem-Solving",
            "Working Style: Detail-focused, organised, proactive, adaptable, quick to learn new systems",
        ],
    },
}


def set_text(paragraph, text: str) -> None:
    first_run = paragraph.runs[0] if paragraph.runs else None
    run_properties = copy.deepcopy(first_run._r.rPr) if first_run is not None else None
    paragraph.clear()
    new_run = paragraph.add_run(text)
    if run_properties is not None:
        new_run._r.insert(0, run_properties)


def find_one(paragraphs, starts_with: str):
    matches = [p for p in paragraphs if p.text.startswith(starts_with)]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph starting with {starts_with!r}; found {len(matches)}")
    return matches[0]


def set_heading_title(paragraph, title: str) -> None:
    if not paragraph.runs:
        raise ValueError(f"Heading has no writable run: {paragraph.text!r}")
    paragraph.runs[0].text = title


def replace_bullet_block(document: Document, anchor_start: str, bullets: list[str], following_heading_start: str) -> None:
    start_index = next(i for i, p in enumerate(document.paragraphs) if p.text.startswith(anchor_start))
    end_index = next(
        i for i, p in enumerate(document.paragraphs[start_index + 1 :], start_index + 1)
        if p.text.startswith(following_heading_start)
    )
    existing = [p for p in document.paragraphs[start_index + 1 : end_index] if p.text.strip()]
    if len(existing) != len(bullets):
        raise ValueError(f"Expected {len(bullets)} bullets after {anchor_start!r}; found {len(existing)}")
    for paragraph, text in zip(existing, bullets):
        set_text(paragraph, text)


def apply_tailoring(file_name: str, content: dict) -> dict:
    target = ROOT / file_name
    backup = BACKUP_ROOT / file_name
    BACKUP_ROOT.mkdir(parents=True, exist_ok=True)
    shutil.copy2(target, backup)

    doc = Document(target)
    paragraphs = doc.paragraphs
    set_text(find_one(paragraphs, "Financial Economics"), content["profile"])

    ideex_start = find_one(paragraphs, "Total Rewards Intern (Data & Reporting)")
    ves_heading = find_one(paragraphs, "Project-Based Research Assistant | University of Copenhagen")
    ideex_start_index = paragraphs.index(ideex_start)
    ves_heading_index = paragraphs.index(ves_heading)
    ideex_paragraphs = [p for p in paragraphs[ideex_start_index + 1 : ves_heading_index] if p.text.strip()]
    if len(ideex_paragraphs) != 4:
        raise ValueError(f"Expected four IDEXX bullets in {file_name}; found {len(ideex_paragraphs)}")
    for paragraph, text in zip(ideex_paragraphs, content["ideex"]):
        set_text(paragraph, text)

    project_heading = find_one(paragraphs, "PROJECTS")
    project_heading_index = paragraphs.index(project_heading)
    ves_paragraphs = [p for p in paragraphs[ves_heading_index + 1 : project_heading_index] if p.text.strip()]
    if len(ves_paragraphs) != 3:
        raise ValueError(f"Expected three VES bullets in {file_name}; found {len(ves_paragraphs)}")
    for paragraph, text in zip(ves_paragraphs, content["ves"]):
        set_text(paragraph, text)

    first_heading = find_one(paragraphs, "FP&A Planning & Performance Model")
    second_heading = find_one(paragraphs, "Customer Profitability Analysis")
    skills_heading = find_one(paragraphs, "SKILLS")
    set_heading_title(first_heading, content["projects"]["first_title"])
    set_heading_title(second_heading, content["projects"]["second_title"])

    first_heading_index = paragraphs.index(first_heading)
    second_heading_index = paragraphs.index(second_heading)
    skills_heading_index = paragraphs.index(skills_heading)
    first_bullets = [p for p in paragraphs[first_heading_index + 1 : second_heading_index] if p.text.strip()]
    second_bullets = [p for p in paragraphs[second_heading_index + 1 : skills_heading_index] if p.text.strip()]
    if len(first_bullets) != 4 or len(second_bullets) != 4:
        raise ValueError(f"Unexpected project bullet count in {file_name}: {len(first_bullets)}, {len(second_bullets)}")
    for paragraph, text in zip(first_bullets, content["projects"]["first_bullets"]):
        set_text(paragraph, text)
    for paragraph, text in zip(second_bullets, content["projects"]["second_bullets"]):
        set_text(paragraph, text)

    skills_end = find_one(paragraphs, "EXTRACURRICULAR ACTIVITIES")
    skills_end_index = paragraphs.index(skills_end)
    skills_paragraphs = [p for p in paragraphs[skills_heading_index + 1 : skills_end_index] if p.text.strip()]
    if len(skills_paragraphs) != 4:
        raise ValueError(f"Expected four skill lines in {file_name}; found {len(skills_paragraphs)}")
    for paragraph, text in zip(skills_paragraphs, content["skills"]):
        set_text(paragraph, text)

    doc.save(target)
    return {"file": str(target), "backup": str(backup), "paragraphs": len(doc.paragraphs)}


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    result = [apply_tailoring(name, content) for name, content in CVS.items()]
    print(json.dumps(result, ensure_ascii=False, indent=2))
