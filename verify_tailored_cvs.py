from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\nguye\OneDrive\Tài liệu\CV & Motivation Letters\New CV")
NAMES = [
    "Minh Nguyen CV - GCS.docx",
    "Minh Nguyen CV - UL Solutions.docx",
    "Minh Nguyen CV - Rosewood Amsterdam.docx",
    "Minh Nguyen CV - McCain Foods.docx",
]


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    for name in NAMES:
        path = ROOT / name
        doc = Document(path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        print(f"=== {name} ===")
        print(f"ZIP_OK={zipfile.ZipFile(path).testzip() is None}")
        print(f"STALE_TEMPLATE={'Financial Economics master’s student with hands-on experience in management reporting' in text}")
        print(f"PLACEHOLDER={any(value in text for value in ('TODO', 'TBD', 'Hunter Douglas'))}")
        for index, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"{index:02d} | {paragraph.text}")
